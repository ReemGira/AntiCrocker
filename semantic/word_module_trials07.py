import math
from docx import Document
from docx.shared import Inches
from docx.shared import Pt

from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

from docx.enum.table import WD_TABLE_ALIGNMENT

from docx.shared import RGBColor

document = Document()
#the following variable is to know if we are at the end of the page or not
currentPage_font_count = 0
#for page number
counter = 1 

def insertHR(paragraph):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)

def max_no_of_choices():
	#this function should return the max number of choices to help in making the table of choices
	#get this number from the DB  test ----------------------------------------------
	return 4

def get_mcq():
    
	#we get the following list from DB	test ----------------------------------------------
	question_header = ["there is very little .... from the factory, so it's nor bad for the environment", \
                    "here is your ticket for the museum , the ticket is ....... for two days." , \
                    "ola spent most of her ..... living on a farm , but she moved to cairo when she was sixteen",\
                    "it ...... that the population if the world is more than seven billion" ,\
                    "nour ... father is a surgeon , is my best friend" , \
                    "i remember things better when i study ....... things such as maps and pictures.",\
                    " the Qsr-ElNile bridge is not ..... the 6th october bridge "]
	return question_header
	
def get_choices():
	#we get the following list from DB	test ----------------------------------------------
	#split choices using ",,__,,"
    question_choices = ["waste,,__,,wave,,__,,wildlife,,__,,weight" , "virtual,,__,,valid,,__,,vinegar,,__,,vapour" , \
                        "child,,__,,childhood,,__,,character,,__,,family" , "believes,,__,,believed,,__,,is believed,,__,, believes" , \
                        "whose,,__,,which,,__,,that,,__,,who" , "wirtual,,__,,seeing,,__,,see,,__,,visual" , \
                        "as long as ,,__,, the long as ,,__,,long as ,,__,, as long"]
    return question_choices

def get_reading():

    list_quotes = """A well-dressed young man entered a big textile shop one evening. He was able to draw the attention of the salesmen who thought him rich and likely to make heavy purchases. He was shown the superior varieties of suit lengths and sarees. But after casually examining them, he kept moving to the next section, where readymade goods were being sold and further on to another section. By then, the salesmen had begun to doubt his intentions and drew the attention of the manager. The manager asked him what exactly he wanted and he replied that he wanted courteous treatment. He explained that he had come to the same shop in casual dress that morning and drawn little attention. His pride was hurt and he wanted to assert himself. He had come in good dress only to get decent treatment, not for getting any textiles. He left without making any purchase."""
	
    list_questions = ["why did the sales man think that the young man would buy lots of clothes ?" , \
                      "what did the young man want when the manager asked him ?", "why did the pride of the young man got hurt ?"\
                      , "what did the sales man do when he doubted about that customer ?" ]	
    list_answers = [len("because he thought he was a rich man and would make heavy purchases") ,\
                 len("he only wanted courteous treatment") , len("he had come to the shop in casual dress , and had little attention.") \
                 , len("he called on his manager.")]
	
    return list_quotes,list_questions,list_answers

def get_quotations():
	list_quotes = "there has been a great argument between the two main politcal groups "
	
	list_questions = ["what did gulliver see inside the kin's palace ?",\
                   "why did the trameksan want to wear high heels on their shoes ?",\
                   "explain if the king'sson was obedient or disobedient to his father."]
	
	list_answers = [60, 20 , 70]
	
	return list_quotes,list_questions,list_answers
	

def get_barcode_no():
		#we get the following list from DB	test ----------------------------------------------
		return 11224444
	
def insertBarcode():
	'''
	pic_par = document.add_paragraph()
	run = pic_par.add_run()
	run.add_picture('barcode03.png', width=Inches(1.0))
	paragraph_format_pic = pic_par.paragraph_format
	paragraph_format_pic.space_before = Pt(0)
	paragraph_format_pic.space_before.pt
	paragraph_format_pic.space_after = Pt(0)
	paragraph_format_pic.space_after.pt
	'''
	barcode = document.add_paragraph()
	
	paragraph_format_barcode = barcode.paragraph_format
	paragraph_format_barcode.space_before = Pt(0)
	paragraph_format_barcode.space_before.pt
	paragraph_format_barcode.space_after = Pt(10)
	paragraph_format_barcode.space_after.pt
	
	barcode_run = barcode.add_run('Code: ')
	barcode_run.bold = True
	font_bar = barcode_run.font
	font_bar.size= Pt(14)
	barcode_run2 = barcode.add_run(str(get_barcode_no()))
	#barcode_run2 = barcode.add_run('   '+str(get_barcode_no()))
	font_bar2 = barcode_run2.font
	font_bar2.size= Pt(12)
	
	#insert page number here at the top of the page
	global counter
	run_page_no = barcode.add_run('\t\t\t\t\t\t\t\t  page '+str(counter))
	font_page_no = run_page_no.font
	font_page_no.size = Pt(12)
	counter += 1

	insertHR(barcode)
	

def myAddPageBreak():
	document.add_page_break()
	insertBarcode()
	global currentPage_font_count
	currentPage_font_count = 0


def my_add_paragraph(text, bold_or_not,change_font, underlined, fontSize=None):
	#global document
	p = document.add_paragraph()
	run = p.add_run(text)
	if(bold_or_not):
		run.bold=True
	if(change_font):
		font = run.font
		font.size = Pt(fontSize)
	if(underlined):
		run.font.underline = True
	return p

def make_answer_rect(last_row,length_answer):
	table = document.add_table(last_row, 2)
	table.style = 'Table Grid'
	
	#write the lines
	my_final_line=""
	my_line = ['_']*105
	my_final_line = my_final_line.join(my_line)
	for i in range(length_answer):
		row = table.rows[i+1].cells
		row[0].text = '\n\n'+my_final_line
		paragraphs = row[0].paragraphs
		for paragraph in paragraphs:
			for run in paragraph.runs:
				font = run.font
				#font.size= Pt(14)
				font.color.rgb = RGBColor(220,220,220)#light gray
				#font.color.rgb = RGBColor(192,192,192)#darker gray=silver#
	
	a = table.cell(0, 0)
	b = table.cell(last_row-1, 1)
	A = a.merge(b)
	
	#document.add_paragraph(str(length_answer))
	
def check_end_of_page(which_part):
	global currentPage_font_count
	max_cnt = 0
	if(which_part):#in the part of mcqs
		max_cnt = 511#530 #519 = 12*5*2 + 19*11 + 19*10 #this value was computed from the real MS word = 12*5*2 + 20*11 + 19*9 
	else :#in the part of essay questions
		max_cnt = 500#512 #this value was computed from the real MS word = (11 * 24) + (12*5) + (10*5*2) +14*2 (for the header only)+12*5(for word answer)
	
	if(currentPage_font_count >= max_cnt):
		myAddPageBreak()
		#document.add_paragraph(str(currentPage_font_count))
		return True
	else:
		#document.add_paragraph(str(currentPage_font_count))
		return False


def make_essay_questions(questions,answer):
	global currentPage_font_count
	i=0
	for str1 in questions:
		#for the answer:
		length_answer = math.floor(answer[i]/30) + 1#math.floor(answer[i]/83) + 1
		last_length = length_answer*3+2
		
		#check if end of page or not
		currentPage_font_count += math.ceil(len(str1)/72)*12+12 +last_length*11 + 10*2 + 12 #10 and 10 for the margins around the question
		if(check_end_of_page(False)):
			currentPage_font_count += math.ceil(len(str1)/72)*12+12 +last_length*11 + 10*2 + 12
		#document.add_paragraph('count= '+str(currentPage_font_count))
		
		#for question
		paragraph = document.add_paragraph('') 
		paragraph_format1 = paragraph.paragraph_format
		paragraph_format1.space_before = Pt(10)
		paragraph_format1.space_before.pt
		paragraph_format1.space_after = Pt(10)
		paragraph_format1.space_after.pt

		#run = paragraph.add_run('Q-'+str(i+1)+':\n'+str1)
		run = paragraph.add_run('Question:\n'+str(i+1)+'- '+str1)
		run.bold = True
		font = run.font
		font.name = 'Calibri'
		font.size = Pt(12)
		
		#for the answer:
		p_answer = my_add_paragraph("Answer:",False,True,False,12)
		paragraph_format2 = p_answer.paragraph_format
		paragraph_format2.space_before = Pt(0)
		paragraph_format2.space_before.pt
		paragraph_format2.space_after = Pt(0)
		paragraph_format2.space_after.pt
		make_answer_rect(last_length,length_answer)
		
		i = i+1



#-------------------------------------------   page format  -----------------------------------------------
#add the QR code at the heading or the top of the page
insertBarcode()

#create the ticket of name and ID after writing 'Exam'
my_add_paragraph("Cairo University",True,True,False,16)
my_add_paragraph("Faculty of Engineering",True,True,False,16)
my_add_paragraph("Computer department",True,True,False,16)

my_add_paragraph("\n\n\t\t\t\t\tExam\n",True,True,False,22)


#rectangle of the ticket name and ID of student
table_merge = document.add_table(6, 2)
table_merge.style = 'Table Grid'

row = table_merge.rows[1].cells
row[0].text = '\nName: '
paragraphs = row[0].paragraphs
for paragraph in paragraphs:
	for run in paragraph.runs:
		font = run.font
		font.size= Pt(14)
			
row = table_merge.rows[3].cells
row[0].text = '\nID: '
paragraphs = row[0].paragraphs
for paragraph in paragraphs:
	for run in paragraph.runs:
		font = run.font
		font.size= Pt(14)

a = table_merge.cell(0, 0)
b = table_merge.cell(5, 1)
A = a.merge(b)


#add notes
my_add_paragraph("\n\n\t\t\t\tImportant Notes\n",True,True,False,18)
my_add_paragraph("\t1. You should write your quartet name",False,True,False,12)
my_add_paragraph("\t2. For essay questions only answers written in the rectangles will be graded",False,True,False,12)
my_add_paragraph("\t3. Your answers shouldn't exceed the space specified below each question",False,True,False,12)
my_add_paragraph("\t4. For Multiple choice questions, only answers in the table will be graded",False,True,False,12)







#-------------------------------------------   MCQ   -----------------------------------------------
myAddPageBreak()

#get the list of mcq questions
pp = my_add_paragraph('Multiple Choice Questions:',True,True,True,14)
paragraph_format1 = pp.paragraph_format
paragraph_format1.space_before = Pt(0)
paragraph_format1.space_before.pt
paragraph_format1.space_after = Pt(14)
paragraph_format1.space_after.pt
currentPage_font_count += 14*2


question_header = get_mcq()
question_choices = get_choices()
#max_choices = max_no_of_choices()
#first create the table for the students to put their mcq answers in
#table = document.add_table(len(question_header)+1, max_choices +1)

all_questions = math.ceil(len(question_header)/23)
last_row_written = 1
while all_questions != 0:
	table = document.add_table(1, 2)
	table.style = 'Table Grid'
	#1. put the header of the table
	header_row = table.rows[0].cells
	header_row[0].text = str(last_row_written)
	header_row[0].width= Inches(0.4)
	header_row[1].width= Inches(1.8)

	paragraphs = header_row[0].paragraphs
	for paragraph in paragraphs:
		for run in paragraph.runs:
			font = run.font
			font.size= Pt(20)
			
	paragraphs = header_row[1].paragraphs
	for paragraph in paragraphs:
		for run in paragraph.runs:
			font = run.font
			font.size= Pt(20)
			

	for x in range(last_row_written-1,len(question_header)-1):
		
		#row_cells[1].text = 'in original loop '+str(all_questions)
		
		if(((x+2)%24) == 0): #means end of table in the current page
			myAddPageBreak()
			last_row_written += 23
			#row_cells[1].text = 'break'+str(last_row_written)+' '+str(all_questions)
			break
			
		row_cells = table.add_row().cells
		row_cells[0].text = str(x+2)
		
		row_cells[0].width= Inches(0.4)
		paragraphs = row_cells[0].paragraphs
		for paragraph in paragraphs:
			for run in paragraph.runs:
				font = run.font
				font.size= Pt(20)
		
		row_cells[1].width= Inches(1.8)
		paragraphs = row_cells[1].paragraphs
		for paragraph in paragraphs:
			for run in paragraph.runs:
				font = run.font
				font.size= Pt(20)
	table.alignment = WD_TABLE_ALIGNMENT.CENTER
	all_questions -= 1
#add another table if the table exceeds the length of a page(make two tables in one page as in bag data exam)
#add a new table in a new page if we had two tables in one page and they already filled the current page


#write the multiple choice questions
myAddPageBreak()
for x in range(len(question_header)):
	
	str1 = question_header[x]
	
	#for choices:
	choices = question_choices[x].split(',,__,,')
	choices_lines_count = 0
	for choice in choices:
		#89 is the max number of charcters written in a line of choices
		choices_lines_count += math.ceil(len(choice)/89)*11 + 9 #10
		
	#check if end of page or not
	currentPage_font_count += math.ceil(len(str1)/72)*12 + 12 + choices_lines_count #12 for the margin after the question
	#document.add_paragraph('count before = '+str(currentPage_font_count))
	if(check_end_of_page(True)):
		currentPage_font_count += math.ceil(len(str1)/72)*12 + 12 + choices_lines_count
	#document.add_paragraph('count= '+str(currentPage_font_count))
	
	
	p = document.add_paragraph() 
	run = p.add_run(str(x+1)+'. '+str1)
	font = run.font
	font.name = 'Times New Roman'
	font.size = Pt(12)
	run.bold = True
	paragraph_format = p.paragraph_format
	paragraph_format.space_after = Pt(12)
	paragraph_format.space_after.pt
	
	#add choices
	
	ch = 'A'
	for choice in choices:
		p1 = document.add_paragraph() 
		run1 = p1.add_run(ch+') '+choice)
		font1 = run1.font
		font1.name = 'Times New Roman'
		font1.size = Pt(11)
		paragraph_format1 = p1.paragraph_format
		paragraph_format1.left_indent
		paragraph_format1.left_indent = Inches(0.3)
		paragraph_format1.space_after = Pt(9)#10)
		paragraph_format1.space_after.pt
		ch = chr(ord(ch) + 1) 


#-------------------------------------------   Reading questions -----------------------------------------------
myAddPageBreak()
#get essay questions, write them and leave appropriate space for answers
pp = my_add_paragraph('Reading Questions:',True,True,True,14)
paragraph_format1 = pp.paragraph_format
paragraph_format1.space_before = Pt(0)
paragraph_format1.space_before.pt
paragraph_format1.space_after = Pt(14)
paragraph_format1.space_after.pt
#document.add_paragraph('count= '+str(currentPage_font_count))
currentPage_font_count += 14*2

Reading_passages,Reading_quesions,Reading_answers = get_reading()


#for paragraph of the reading
paragraph = my_add_paragraph(Reading_passages+'\n',False,True,False,13)
paragraph_format1 = paragraph.paragraph_format
paragraph_format1.space_before = Pt(10)
paragraph_format1.space_before.pt
paragraph_format1.space_after = Pt(10)
paragraph_format1.space_after.pt

currentPage_font_count += (math.ceil(len(Reading_passages)/79))*13+13 + 40 #10*4 for the margins around the paragraph and around the endline after the paragraph

make_essay_questions(Reading_quesions,Reading_answers)



#-------------------------------------------   Essay questions -----------------------------------------------
myAddPageBreak()
#get essay questions, write them and leave appropriate space for answers
pp = my_add_paragraph('Story Questions:',True,True,True,14)
paragraph_format1 = pp.paragraph_format
paragraph_format1.space_before = Pt(0)
paragraph_format1.space_before.pt
paragraph_format1.space_after = Pt(14)
paragraph_format1.space_after.pt
#document.add_paragraph('count= '+str(currentPage_font_count))
currentPage_font_count += 14*2
#document.add_paragraph('now count= '+str(currentPage_font_count))

#who said that to whom question --> quotations question:
quotation, quotation_questions, quotation_answers = get_quotations()

#writing the quotation:
paragraph = my_add_paragraph("for the following quote, answer the below questions:",True,True,True,13)
paragraph_format1 = paragraph.paragraph_format
paragraph_format1.space_before = Pt(10)
paragraph_format1.space_before.pt
paragraph_format1.space_after = Pt(10)
paragraph_format1.space_after.pt

my_add_paragraph("\""+quotation+"\"",True,True,False,13)
currentPage_font_count += (math.ceil(len(quotation)/79))*13 + 13 + 10*4

make_essay_questions(quotation_questions, quotation_answers)
'''
#test ----------------------------------------------
#not quotation questions:

questions = ["Write your name!","what is your age?","are you there?","I was sailing in the sea","what is your age?","Write your name!","what is your age?","are you there?"]
answer = [110, 3,50,10,3,110, 3,50]

myAddPageBreak()

#for other story questions:
paragraph = my_add_paragraph("Answer the following questions:",True,True,True,13)
paragraph_format1 = paragraph.paragraph_format
paragraph_format1.space_before = Pt(10)
paragraph_format1.space_before.pt
paragraph_format1.space_after = Pt(10)
paragraph_format1.space_after.pt

currentPage_font_count += 13 + 20 #20 = 10*2 for the margins 


make_essay_questions(questions,answer)
'''		


document.save('mydemo.docx')
